from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from config.database import (
    collection_name,
    collection_knowledge,
    collection_questions,
    collection_knowledge_for_pinecone,
)
from transformers import (
    WhisperProcessor,
    WhisperForConditionalGeneration,
    SpeechT5Processor,
    SpeechT5ForTextToSpeech,
    SpeechT5HifiGan,
    pipeline,
)
from models.text import textClass
from models.question import QuestionClass
import librosa, io
import torch
import soundfile as sf
from datasets import load_dataset
import os
from docx import Document
from datetime import datetime
from functions.questions_answer import (
    extract_bracketed_sentences,
    separate_brackets,
)
import base64
from typing import List
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import Pinecone
from langchain.llms import OpenAI
import pinecone
from langchain.chains.question_answering import load_qa_chain

router = APIRouter()


# endpoint for transcribing audio to text
@router.post("/audio")
async def upload_audio(audio_file: UploadFile = File(...)):
    # baca audio file
    content = await audio_file.read()

    # Simpan file audio ke MongoDB
    file_id = collection_name.insert_one({"content": content}).inserted_id

    # load model and processor
    processor = WhisperProcessor.from_pretrained(os.getenv("SMALL_WHISPER"))
    model = WhisperForConditionalGeneration.from_pretrained(os.getenv("SMALL_WHISPER"))
    forced_decoder_ids = processor.get_decoder_prompt_ids(
        language="english", task="transcribe"
    )
    model.config.forced_decoder_ids = forced_decoder_ids

    # Proses audio untuk pengenalan teks
    waveform, sampling_rate = librosa.load(io.BytesIO(content), sr=16000)
    input_features = processor(
        waveform, sampling_rate=sampling_rate, return_tensors="pt"
    ).input_features
    predicted_ids = model.generate(input_features)
    transcription = processor.batch_decode(predicted_ids, skip_special_tokens=True)

    if transcription:
        return {"file_id": str(file_id), "transcription": transcription[0]}

    else:
        raise HTTPException(status_code=500, detail="Speech recognition failed.")


# endpoint for text to speech
@router.post("/text")
async def text_to_speech(request: textClass):
    # Load the pre-trained models
    processor = SpeechT5Processor.from_pretrained(os.getenv("TEXT_TO_SPEECH"))
    model = SpeechT5ForTextToSpeech.from_pretrained(os.getenv("TEXT_TO_SPEECH"))
    vocoder = SpeechT5HifiGan.from_pretrained(os.getenv("TEXT_TO_SPEECH_HIFIGAN"))

    # Load xvector containing speaker's voice characteristics from a dataset
    embeddings_dataset = load_dataset(
        "Matthijs/cmu-arctic-xvectors", split="validation"
    )
    speaker_embeddings = torch.tensor(embeddings_dataset[7306]["xvector"]).unsqueeze(0)

    # Get the text input from the user
    input_text = request.text

    # Validate input
    if not input_text.strip():
        raise HTTPException(
            status_code=400, detail="Invalid input: Text cannot be empty."
        )

    # Generate audio
    inputs = processor(text=input_text, return_tensors="pt")
    speech = model.generate_speech(
        inputs.input_ids, speaker_embeddings, vocoder=vocoder
    )

    # Save the audio in memory
    audio_buffer = io.BytesIO()
    sf.write(audio_buffer, speech.numpy(), samplerate=16000, format="wav")
    audio_buffer.seek(0)

    # Encode the audio data in base64
    audio_base64 = base64.b64encode(audio_buffer.read()).decode("utf-8")
    print(audio_base64)
    return JSONResponse(content={"base64_audio": audio_base64})


# endpoint for updating knowledge to database MongoDB
@router.post("/knowledge")
async def update_knowledge(
    knowledge_file: UploadFile = File(...),
    category_context: str = Form(...),
):
    try:
        document = Document(knowledge_file.file)
        knowledge_text = ""
        for paragraph in document.paragraphs:
            knowledge_text += paragraph.text + " "

        knowledge_text = separate_brackets(knowledge_text)

        # Check knowledge based on category, then remove it when it exists
        existing_knowledge = collection_knowledge.find_one(
            {"category_context": category_context}
        )
        if existing_knowledge:
            collection_knowledge.delete_one({"category_context": category_context})
            print(
                "Knowledge with category {} has been deleted".format(category_context)
            )

        # add new knowledge
        collection_knowledge.insert_one(
            {
                "category_context": category_context,
                "text": knowledge_text,
            }
        )
        print("knowledge with category {} added".format(category_context))
        return {
            "message": "Knowledge context added successfully",
            "category_context": category_context,
            "knowledge_text": knowledge_text,
        }

    except Exception as e:
        raise HTTPException(detail=e.args[0], status_code=400)


# endpoint for question answer with huggingface model
@router.post("/question")
async def question_answer(question: QuestionClass):
    try:
        # Setup Question Answering model
        qa_pipeline = pipeline(
            "question-answering", model="timpal0l/mdeberta-v3-base-squad2"
        )

        # Specify the category context directly
        category_context = os.getenv("ENGLISH_CONTEXT")

        context = ""
        for doc in collection_knowledge.find(
            {"category_context": category_context}, {"text": 1}
        ):
            context += doc["text"] + " "
            print(context)

        result = qa_pipeline({"context": context, "question": question.text})
        answer = result["answer"]
        confidence = result["score"]  # Mendapatkan tingkat confidence
        print("Level confidence: " + str(confidence))

        if confidence > 0.005:
            target_sentence = answer
            target_index = result["start"]

            matching_sentences = extract_bracketed_sentences(
                context, target_sentence, target_index
            )

            if len(matching_sentences) > 0:
                print("Matching bracketed sentences:")
                for sentence in matching_sentences:
                    print(sentence.strip())

            else:
                print("No matching bracketed sentences found.")

            # Save question to database
            question_doc = {
                "text": question.text,
                "createdAt": datetime.now(),
                "category_context": category_context,
            }

            question_text = collection_questions.insert_one(question_doc)
            inserted_id = str(question_text.inserted_id)

            return {
                "relevant_answer": matching_sentences,
                "questions_saved": {
                    "question_id": inserted_id,
                    "question_text": question.text,
                    "createdAt": question.createdAt,
                    "category_context": category_context,
                },
            }
        else:
            return "Level confidence is to low"

    except Exception as e:
        raise HTTPException(status_code=400, detail=e.args[0])


# endpoint for update knowledge to pinecone
@router.post("/knowledge_pinecone")
async def update_knowledge_pinecone(
    knowledge_file: UploadFile = File(...),
    category_context: str = Form(...),
):
    try:
        # read the document
        document = Document(knowledge_file.file)
        knowledge_text = ""
        for paragraph in document.paragraphs:
            knowledge_text += paragraph.text + " "

        # check existing knowledge in the database mongoDB
        existing_knowledge = collection_knowledge_for_pinecone.find_one(
            {"category_context": category_context}
        )

        # if exist, delete the knowledge
        if existing_knowledge:
            collection_knowledge_for_pinecone.delete_one(
                {"category_context": category_context}
            )
            print("knowledge with category {} is deleted".format(category_context))

        # add new knowledge
        collection_knowledge_for_pinecone.insert_one(
            {
                "category_context": category_context,
                "text": knowledge_text,
            }
        )
        print("knowledge with category {} added".format(category_context))

        # split the text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000, chunk_overlap=200, length_function=len
        )

        # convert the chunks into documents
        text_splitted = text_splitter.create_documents([knowledge_text])
        print(len(text_splitted))

        # initialize vector store pinecone
        pinecone.init(
            api_key=os.getenv("PINECONE_API_KEY"),
            environment=os.getenv("PINECONE_ENV"),
        )
        index_name = "chatbot-bni-direct"

        # load OpenAI Embeddings model
        embeddings = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
        if index_name not in pinecone.list_indexes():
            print("Index does not exist", index_name)

        book_docsearch = Pinecone.from_texts(
            [t.page_content for t in text_splitted], embeddings, index_name=index_name
        )

        return {
            "knowledge_text": knowledge_text,
            "text_splitted": text_splitted,
            "category_context": category_context,
        }

    except Exception as e:
        raise HTTPException(status_code=400, detail=e.args[0])


# endpoint for answering using llm openAI
@router.post("/answer")
async def answering(question: QuestionClass):
    llm = OpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"))
    print("load llm berhasil")
    query = question.text
    index_name = "chatbot-bni-direct"
    embeddings = OpenAIEmbeddings(openai_api_key=os.getenv("OPENAI_API_KEY"))
    print("load embeddings berhasil")
    pinecone.init(
        api_key=os.getenv("PINECONE_API_KEY"),
        environment=os.getenv("PINECONE_ENV"),
    )
    docsearch = Pinecone.from_existing_index(index_name, embeddings)
    print("load docsearch berhasil")

    docs = docsearch.similarity_search(query)
    print("mencari similarity berhasil")
    chain = load_qa_chain(llm, chain_type="stuff")
    result = chain.run(input_documents=docs, question=query)

    return {"answer": result}
