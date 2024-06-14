import os, sys
import pytest
from unittest.mock import patch, MagicMock
from fastapi import UploadFile
from io import BytesIO
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from functions.knowledge import add_knowledge
from docx import Document

def create_test_docx():
    doc = Document()
    doc.add_paragraph("This is a test document.")
    with BytesIO() as f:
        doc.save(f)
        f.seek(0)
        return f.read()

@pytest.fixture
def mock_upload_file():
    file_content = create_test_docx()
    file = BytesIO(file_content)
    upload_file = UploadFile(filename="test.docx", file=file)
    return upload_file

@patch("functions.knowledge.Document")
@patch("functions.knowledge.SentenceTransformerEmbeddings")
@patch("functions.knowledge.PGVector")
@patch.dict(os.environ, {"EMBEDDING_MODEL": "model_name", "PGVECTOR_CONNECTION_STRING": "connection_string", "COLLECTION_NAME": "collection_name"})
def test_add_knowledge(mock_pgvector, mock_embeddings, mock_document, mock_upload_file):

    # Mock Document to return a custom document
    mock_doc_instance = MagicMock()
    mock_doc_instance.paragraphs = [MagicMock(text="This is a test document.")]
    mock_document.return_value = mock_doc_instance

    # Mock the embeddings instance
    mock_embeddings_instance = MagicMock()
    mock_embeddings.return_value = mock_embeddings_instance
    mock_pgvector.from_documents.return_value = MagicMock()

    # Call the function
    response = add_knowledge(mock_upload_file)

    # Assertions
    assert response == {"message": "knowledge has been added successfully"}
    mock_document.assert_called_once_with(mock_upload_file.file)
    mock_embeddings.assert_called_once_with(model_name="model_name")

    # Verifying the structure of documents instead of using MagicMock
    args, kwargs = mock_pgvector.from_documents.call_args
    assert kwargs["embedding"] == mock_embeddings_instance
    assert len(kwargs["documents"]) == 1
    assert kwargs["documents"][0].page_content == "This is a test document."
    assert kwargs["collection_name"] == "collection_name"
    assert kwargs["connection_string"] == "connection_string"
